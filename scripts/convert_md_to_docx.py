# -*- coding: utf-8 -*-
"""二等学科対策問題集のMarkdownをWord(.docx)へ変換するスクリプト。
使い方: python convert_md_to_docx.py <input.md> <output.docx>
"""
import sys, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_runs_with_bold(paragraph, text):
    """**bold** 記法を解釈してrunを追加する。"""
    for i, part in enumerate(re.split(r'(\*\*.*?\*\*)', text)):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def convert(in_path, out_path):
    lines = open(in_path, encoding='utf-8').read().split('\n')
    doc = Document()
    # 既定フォント（日本語対応）
    style = doc.styles['Normal']
    style.font.name = 'Yu Gothic'
    style.font.size = Pt(10.5)

    for raw in lines:
        line = raw.rstrip()
        if line.strip() == '' or line.strip() == '---':
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith('> '):
            p = doc.add_paragraph()
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif re.match(r'^\*\(.*\)\*$', line.strip()):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().strip('*'))
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif re.match(r'^\|', line):
            # 表は本問題集では未使用のためそのまま段落化
            doc.add_paragraph(line)
        else:
            p = doc.add_paragraph()
            add_runs_with_bold(p, line)

    doc.save(out_path)
    print('saved:', out_path)


if __name__ == '__main__':
    convert(sys.argv[1], sys.argv[2])
